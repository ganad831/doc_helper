"""Integration tests for WordDocumentAdapter."""

import tempfile
from pathlib import Path

import pytest
from docx import Document

from doc_helper.domain.common.result import Failure, Success
from doc_helper.domain.document.document_format import DocumentFormat
from doc_helper.infrastructure.document.word_document_adapter import (
    WordDocumentAdapter,
)


class TestWordDocumentAdapter:
    """Integration tests for WordDocumentAdapter."""

    @pytest.fixture
    def adapter(self) -> WordDocumentAdapter:
        """Create adapter instance."""
        return WordDocumentAdapter()

    @pytest.fixture
    def template_file(self, tmp_path: Path) -> Path:
        """Create a simple Word template with content controls."""
        template_path = tmp_path / "template.docx"

        # Create a simple Word document
        doc = Document()
        doc.add_paragraph("Test Template")
        doc.add_paragraph("Project Name: ")
        doc.add_paragraph("Date: ")
        doc.save(str(template_path))

        return template_path

    def test_adapter_format(self, adapter: WordDocumentAdapter) -> None:
        """Adapter should report WORD format."""
        assert adapter.format == DocumentFormat.WORD

    def test_validate_template_success(
        self, adapter: WordDocumentAdapter, template_file: Path
    ) -> None:
        """validate_template should succeed for valid template."""
        result = adapter.validate_template(template_file)
        assert isinstance(result, Success)

    def test_validate_template_not_found(
        self, adapter: WordDocumentAdapter
    ) -> None:
        """validate_template should fail for non-existent template."""
        result = adapter.validate_template("nonexistent.docx")
        assert isinstance(result, Failure)
        assert "not found" in result.error.lower()

    def test_validate_template_invalid_file(
        self, adapter: WordDocumentAdapter, tmp_path: Path
    ) -> None:
        """validate_template should fail for invalid Word file."""
        # Create a non-Word file
        invalid_file = tmp_path / "invalid.docx"
        invalid_file.write_text("Not a Word document")

        result = adapter.validate_template(invalid_file)
        assert isinstance(result, Failure)
        assert "invalid" in result.error.lower()

    def test_generate_document_success(
        self, adapter: WordDocumentAdapter, template_file: Path, tmp_path: Path
    ) -> None:
        """generate should create output document."""
        output_path = tmp_path / "output.docx"
        field_values = {"project_name": "Test Project", "date": "2024-01-15"}

        result = adapter.generate(
            template_path=template_file,
            output_path=output_path,
            field_values=field_values,
        )

        assert isinstance(result, Success)
        assert output_path.exists()

    def test_generate_document_creates_parent_directories(
        self, adapter: WordDocumentAdapter, template_file: Path, tmp_path: Path
    ) -> None:
        """generate should create parent directories if needed."""
        output_path = tmp_path / "subdir" / "output.docx"
        field_values = {}

        result = adapter.generate(
            template_path=template_file,
            output_path=output_path,
            field_values=field_values,
        )

        assert isinstance(result, Success)
        assert output_path.exists()
        assert output_path.parent.exists()

    def test_generate_with_invalid_template(
        self, adapter: WordDocumentAdapter, tmp_path: Path
    ) -> None:
        """generate should fail with invalid template."""
        output_path = tmp_path / "output.docx"
        field_values = {}

        result = adapter.generate(
            template_path="nonexistent.docx",
            output_path=output_path,
            field_values=field_values,
        )

        assert isinstance(result, Failure)

    def test_generated_document_is_valid(
        self, adapter: WordDocumentAdapter, template_file: Path, tmp_path: Path
    ) -> None:
        """Generated document should be a valid Word document."""
        output_path = tmp_path / "output.docx"
        field_values = {"test_field": "test_value"}

        adapter.generate(
            template_path=template_file,
            output_path=output_path,
            field_values=field_values,
        )

        # Should be able to open the generated document
        doc = Document(str(output_path))
        assert doc is not None
        assert len(doc.paragraphs) > 0
